import importlib.util
import tempfile
import unittest
from pathlib import Path

from docx import Document


def load_build_report_docx_module():
    module_path = Path(__file__).resolve().parent.parent / "scripts" / "build_report_docx.py"
    spec = importlib.util.spec_from_file_location("build_report_docx", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


class BuildReportDocxTests(unittest.TestCase):
    def test_build_docx_preserves_tables_wrapped_in_tex4ht_divs(self) -> None:
        module = load_build_report_docx_module()

        html = """\
<html>
  <body>
    <h3>Evidence from BenchResults</h3>
    <div class="table">
      <p class="indent"><a id="x1-43001r4"></a></p>
      <hr class="float" />
      <div class="float">
        <div class="caption"><span class="id">Table 4: </span><span class="content">Pass-rate snapshot from the consolidated BenchResults artefacts</span></div>
        <div class="tabular">
          <table id="TBL-5" class="tabular">
            <tr>
              <td><p class="noindent">Model</p></td>
              <td><p class="noindent">HumanEval</p></td>
            </tr>
            <tr>
              <td><p class="noindent">qwen2.5-coder:7b</p></td>
              <td><p class="noindent">0.91</p></td>
            </tr>
          </table>
        </div>
      </div>
      <hr class="endfloat" />
    </div>
  </body>
</html>
"""

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            html_path = tmpdir_path / "report.html"
            docx_path = tmpdir_path / "report.docx"
            html_path.write_text(html, encoding="utf-8")

            module.build_docx(html_path, docx_path)

            document = Document(docx_path)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(document.tables[0].cell(0, 0).text, "Model")
            self.assertEqual(document.tables[0].cell(1, 0).text, "qwen2.5-coder:7b")

            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            self.assertIn(
                "Table 4: Pass-rate snapshot from the consolidated BenchResults artefacts",
                paragraphs,
            )


if __name__ == "__main__":
    unittest.main()
